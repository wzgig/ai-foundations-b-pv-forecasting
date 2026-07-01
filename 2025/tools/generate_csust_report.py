# -*- coding: utf-8 -*-
"""Generate the course report in a CSUST thesis-sample-like layout.

The source is ``2025/05_delivery/项目主报告_终稿.md``.  The generated DOCX is
then opened by Word in the release workflow to update the table of contents
and export PDF.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "05_delivery"
SOURCE = REPORT_DIR / "项目主报告_终稿.md"
OUTPUT_DOCX = REPORT_DIR / "项目主报告_终稿.docx"
TRACKED_LOGO = ROOT / "03_figures" / "paper_assets" / "csust_header_logo.jpeg"
EXTRACTED_TEMPLATE_LOGO = ROOT.parent / "tmp" / "template_review" / "media" / "image2.jpeg"
TEMPLATE_LOGO = TRACKED_LOGO if TRACKED_LOGO.exists() else EXTRACTED_TEMPLATE_LOGO
TITLE = "光伏电站日前计划与功率预测工作台项目主报告"
EN_TITLE = "PV STATION DAY-AHEAD PLANNING AND POWER FORECASTING WORKBENCH"


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 12, bold: bool | None = None) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, first_line: float = 24, line_spacing: float = 18, justify: bool = True) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(first_line)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_text(paragraph, text: str, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 12, bold: bool | None = None) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)


def add_centered(doc: Document, text: str, east_asia: str, latin: str = "Times New Roman", size: float = 15, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_text(p, text, east_asia=east_asia, latin=latin, size=size, bold=bold)


def should_left_align(text: str) -> bool:
    technical_markers = ("/", "\\", "_", "GitHub", "Streamlit", "HTTP", "Python", "NWP", "LMD", "outputs")
    ascii_count = sum(1 for char in text if ord(char) < 128 and not char.isspace())
    return ascii_count >= 12 and any(marker in text for marker in technical_markers)


def add_body_paragraph(doc: Document, text: str, english: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=24, line_spacing=18, justify=english or not should_left_align(text))
    if english:
        add_text(p, text, east_asia="宋体", latin="Times New Roman", size=12)
    else:
        add_text(p, text, east_asia="宋体", latin="Times New Roman", size=12)


def add_heading(doc: Document, text: str, level: int) -> None:
    style_name = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(18)
        size = 15
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        size = 14
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.space_before = Pt(7.8)
        p.paragraph_format.space_after = Pt(7.8)
        size = 12
    add_text(p, text, east_asia="黑体", latin="Times New Roman", size=size)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def set_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "6", "color": "000000"},
                bottom={"val": "single", "sz": "6", "color": "000000"},
                left={"val": "nil"},
                right={"val": "nil"},
            )


def infer_table_caption(rows: list[list[str]]) -> str:
    header = rows[0] if rows else []
    if header[:3] == ["角色", "主要职责", "交付成果"]:
        return "团队角色分工表"
    if header[:3] == ["模块", "主要路径", "作用"]:
        return "系统模块划分表"
    if header[:2] == ["链路", "模型或口径"]:
        return "预测链路关键指标对比表"
    if header[:3] == ["测试项", "测试命令或入口", "预期结果"]:
        return "功能与稳定性测试表"
    if header[:2] == ["对比项", "结果"]:
        return "结果提升对比表"
    if header[:3] == ["稳定性场景", "设计措施", "效果"]:
        return "稳定性设计表"
    return "项目数据表"


def add_table_caption(doc: Document, table_number: int, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    add_text(p, f"表 {table_number} {caption}", east_asia="宋体", latin="Times New Roman", size=10.5)
    keep_with_next(p)


def add_markdown_table(doc: Document, rows: list[list[str]], table_number: int) -> None:
    if not rows:
        return
    if table_number == 1:
        doc.add_page_break()
    add_table_caption(doc, table_number, infer_table_caption(rows))
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for word_row in table.rows:
        keep_row_together(word_row)
    for r_index, row in enumerate(rows):
        for c_index, cell_text in enumerate(row):
            cell = table.cell(r_index, c_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(18)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            add_text(p, cell_text, east_asia="宋体", latin="Times New Roman", size=10.5, bold=(r_index == 0))
    set_table_borders(table)


def add_image(doc: Document, source_line: str, base: Path, figure_number: int) -> None:
    match = re.match(r"!\[(.*?)\]\((.*?)\)", source_line.strip())
    if not match:
        return
    caption, raw_path = match.groups()
    image_path = (base / raw_path).resolve()
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    keep_with_next(p)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(13.2))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    caption_p.paragraph_format.line_spacing = Pt(12)
    caption_text = caption if caption.startswith("图") else f"图 {figure_number} {caption}"
    add_text(caption_p, caption_text, east_asia="宋体", latin="Times New Roman", size=10.5)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    if placeholder:
        run._r.append(OxmlElement("w:t"))
        run._r[-1].text = placeholder

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_toc(doc: Document) -> None:
    add_centered(doc, "目  录", east_asia="黑体", size=16)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u', "更新目录")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for idx, size in [(1, 15), (2, 14), (3, 12)]:
        style = doc.styles[f"Heading {idx}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor(0, 0, 0)

    for style_name, size in [("toc 1", 11), ("toc 2", 10.5), ("toc 3", 10)]:
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(14)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def configure_section(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    header = section.header
    header.is_linked_to_previous = False
    header_table = header.add_table(rows=1, cols=2, width=Cm(16))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = header_table.rows[0].cells
    if TEMPLATE_LOGO.exists():
        left.paragraphs[0].add_run().add_picture(str(TEMPLATE_LOGO), width=Cm(3.6))
    else:
        add_text(left.paragraphs[0], "长沙理工大学", east_asia="黑体", size=12)
    right_p = right.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(right_p, TITLE, east_asia="宋体", size=9)
    for cell in header_table.rows[0].cells:
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": "000000"}, top={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "第 ", east_asia="宋体", size=9)
    add_field(p, "PAGE", "1")
    add_text(p, " 页 共 ", east_asia="宋体", size=9)
    add_field(p, "NUMPAGES", "1")
    add_text(p, " 页", east_asia="宋体", size=9)


def strip_markup(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    return text.strip()


def split_table_row(line: str) -> list[str]:
    return [strip_markup(part.strip()) for part in line.strip().strip("|").split("|")]


def parse_sections(markdown: str) -> tuple[list[str], list[str], list[str]]:
    zh = re.search(r"## 摘要\s+(.*?)\n\*\*关键词：\*\*\s*(.*?)\n", markdown, flags=re.S)
    en = re.search(r"## ABSTRACT\s+(.*?)\n\*\*Key words:\*\*\s*(.*?)\n", markdown, flags=re.S)
    body = re.search(r"## 1\. 项目概述\s+(.*)", markdown, flags=re.S)

    zh_parts = []
    if zh:
        zh_text = zh.group(1).strip()
        zh_parts = [strip_markup(p) for p in zh_text.split("\n\n") if p.strip()]
        zh_parts.append("关键词：" + strip_markup(zh.group(2)))

    en_parts = []
    if en:
        en_text = en.group(1).strip()
        en_parts = [strip_markup(p) for p in en_text.split("\n\n") if p.strip() and not p.strip().startswith("**PV")]
        en_parts.append("Key words: " + strip_markup(en.group(2)))

    body_lines = []
    if body:
        body_lines = ["## 1. 项目概述", *body.group(1).splitlines()]
    return zh_parts, en_parts, body_lines


def add_body_from_markdown(doc: Document, lines: list[str], base: Path) -> None:
    i = 0
    in_code = False
    code_lines: list[str] = []
    figure_number = 1
    table_number = 1
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            if in_code:
                for code in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(24)
                    add_text(p, code, east_asia="宋体", latin="Consolas", size=10)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("!["):
            add_image(doc, stripped, base, figure_number)
            figure_number += 1
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}", lines[i + 1]):
            table_lines = [stripped]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, [split_table_row(row) for row in table_lines], table_number)
            table_number += 1
            continue

        if stripped.startswith("## "):
            add_heading(doc, strip_markup(stripped[3:]), 1)
        elif stripped.startswith("### "):
            add_heading(doc, strip_markup(stripped[4:]), 2)
        elif stripped.startswith("#### "):
            add_heading(doc, strip_markup(stripped[5:]), 3)
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=24, line_spacing=18, justify=not should_left_align(stripped))
            add_text(p, strip_markup(stripped), east_asia="宋体", size=12)
        elif stripped.startswith("- "):
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=24, line_spacing=18, justify=not should_left_align(stripped))
            add_text(p, strip_markup(stripped), east_asia="宋体", size=12)
        else:
            add_body_paragraph(doc, strip_markup(stripped))
        i += 1


def build_report() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    zh_parts, en_parts, body_lines = parse_sections(markdown)

    doc = Document()
    configure_styles(doc)
    configure_section(doc)

    add_centered(doc, TITLE, east_asia="黑体", size=16)
    doc.add_paragraph()
    add_centered(doc, "摘要", east_asia="黑体", size=15)
    doc.add_paragraph()
    for part in zh_parts:
        add_body_paragraph(doc, part)

    doc.add_page_break()
    add_centered(doc, EN_TITLE, east_asia="Times New Roman", latin="Times New Roman", size=16, bold=True)
    doc.add_paragraph()
    add_centered(doc, "ABSTRACT", east_asia="Times New Roman", latin="Times New Roman", size=15, bold=True)
    doc.add_paragraph()
    for part in en_parts:
        add_body_paragraph(doc, part, english=True)

    doc.add_page_break()
    add_toc(doc)
    doc.add_page_break()
    add_body_from_markdown(doc, body_lines, SOURCE.parent)

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_report()
